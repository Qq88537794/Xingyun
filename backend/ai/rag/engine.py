"""
RAG引擎
整合分块、嵌入、检索、重排序等功能的统一入口
"""

import logging
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
import os

from .chunker import TextChunker, ChunkingStrategy, TextChunk
from .embedding import EmbeddingService, init_embedding_service
from .vector_store import VectorStore, init_vector_store, DistanceMetric
from .retriever import HybridRetriever, RetrievalResult
from .reranker import Reranker, RerankerConfig, RerankerModel

logger = logging.getLogger(__name__)


@dataclass
class RAGConfig:
    """RAG配置"""
    # 分块配置
    chunk_size: int = 500
    chunk_overlap: int = 50
    chunking_strategy: ChunkingStrategy = ChunkingStrategy.RECURSIVE
    
    # 嵌入配置
    embedding_provider: str = "zhipu"
    embedding_model: Optional[str] = None
    
    # 向量存储配置
    persist_directory: Optional[str] = None
    distance_metric: DistanceMetric = DistanceMetric.COSINE
    
    # 检索配置
    vector_weight: float = 0.7
    keyword_weight: float = 0.3
    use_hybrid: bool = True
    
    # 重排序配置
    rerank_enabled: bool = True
    rerank_model: RerankerModel = RerankerModel.SIMPLE
    
    # 结果配置
    top_k: int = 5
    min_score: float = 0.3


@dataclass
class IndexedDocument:
    """已索引的文档信息"""
    doc_id: str
    filename: str
    chunk_count: int
    total_tokens: int
    metadata: Dict[str, Any] = field(default_factory=dict)


class RAGEngine:
    """
    RAG引擎
    提供完整的文档索引和检索功能
    """
    
    def __init__(
        self,
        config: Optional[RAGConfig] = None,
        embedding_api_key: Optional[str] = None,
        llm_client: Optional[Any] = None
    ):
        """
        初始化RAG引擎
        
        Args:
            config: RAG配置
            embedding_api_key: 嵌入模型API密钥
            llm_client: LLM客户端(用于LLM重排序)
        """
        self.config = config or RAGConfig()
        self._llm_client = llm_client
        
        # 初始化组件
        self._chunker = TextChunker(
            chunk_size=self.config.chunk_size,
            chunk_overlap=self.config.chunk_overlap,
            strategy=self.config.chunking_strategy
        )
        
        self._embedding_service = EmbeddingService(
            provider=self.config.embedding_provider
        )
        if embedding_api_key:
            self._embedding_service.configure(
                api_key=embedding_api_key,
                model_name=self.config.embedding_model
            )
        
        self._vector_store = VectorStore(
            persist_directory=self.config.persist_directory,
            distance_metric=self.config.distance_metric
        )
        
        self._retriever = HybridRetriever(
            vector_store=self._vector_store,
            embedding_service=self._embedding_service,
            vector_weight=self.config.vector_weight,
            keyword_weight=self.config.keyword_weight
        )
        
        self._reranker = Reranker(
            config=RerankerConfig(
                model=self.config.rerank_model,
                top_k=self.config.top_k,
                min_score=self.config.min_score
            ),
            llm_client=llm_client
        )
        
        # 已索引文档记录
        self._indexed_docs: Dict[str, IndexedDocument] = {}
    
    def configure_embedding(
        self,
        api_key: str,
        provider: Optional[str] = None,
        model_name: Optional[str] = None
    ):
        """配置嵌入服务"""
        if provider:
            self._embedding_service = EmbeddingService(provider=provider)
        
        self._embedding_service.configure(
            api_key=api_key,
            model_name=model_name
        )
        
        # 重新创建检索器
        self._retriever = HybridRetriever(
            vector_store=self._vector_store,
            embedding_service=self._embedding_service,
            vector_weight=self.config.vector_weight,
            keyword_weight=self.config.keyword_weight
        )
    
    def index_text(
        self,
        text: str,
        doc_id: str,
        metadata: Optional[Dict[str, Any]] = None,
        collection_name: Optional[str] = None
    ) -> IndexedDocument:
        """
        索引文本
        
        Args:
            text: 文本内容
            doc_id: 文档ID
            metadata: 元数据
            collection_name: 集合名称
        
        Returns:
            IndexedDocument
        """
        metadata = metadata or {}
        
        # 分块
        chunks = self._chunker.chunk_text(
            text=text,
            metadata=metadata,
            doc_id=doc_id
        )
        
        if not chunks:
            logger.warning(f"文档 {doc_id} 分块结果为空")
            return IndexedDocument(
                doc_id=doc_id,
                filename=metadata.get('filename', doc_id),
                chunk_count=0,
                total_tokens=0,
                metadata=metadata
            )
        
        # 准备索引数据
        documents = [chunk.content for chunk in chunks]
        chunk_ids = [chunk.id for chunk in chunks]
        chunk_metadatas = [
            {
                **metadata,
                'doc_id': doc_id,
                'chunk_index': chunk.chunk_index,
                'start_index': chunk.start_index,
                'end_index': chunk.end_index
            }
            for chunk in chunks
        ]
        
        # 索引到检索器
        self._retriever.index_documents(
            documents=documents,
            metadatas=chunk_metadatas,
            ids=chunk_ids,
            collection_name=collection_name
        )
        
        # 记录文档信息
        total_tokens = sum(chunk.token_count for chunk in chunks)
        indexed_doc = IndexedDocument(
            doc_id=doc_id,
            filename=metadata.get('filename', doc_id),
            chunk_count=len(chunks),
            total_tokens=total_tokens,
            metadata=metadata
        )
        self._indexed_docs[doc_id] = indexed_doc
        
        logger.info(f"文档 {doc_id} 索引完成: {len(chunks)} 块, {total_tokens} tokens")
        return indexed_doc
    
    def index_file(
        self,
        file_path: str,
        doc_id: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None,
        collection_name: Optional[str] = None
    ) -> IndexedDocument:
        """
        索引文件
        
        Args:
            file_path: 文件路径
            doc_id: 文档ID(默认使用文件名)
            metadata: 额外元数据
            collection_name: 集合名称
        
        Returns:
            IndexedDocument
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        filename = os.path.basename(file_path)
        doc_id = doc_id or filename
        
        # 读取文件内容
        text = self._read_file(file_path)
        
        # 构建元数据
        file_metadata = {
            'filename': filename,
            'file_path': file_path,
            'file_size': os.path.getsize(file_path),
            **(metadata or {})
        }
        
        return self.index_text(
            text=text,
            doc_id=doc_id,
            metadata=file_metadata,
            collection_name=collection_name
        )
    
    def _read_file(self, file_path: str) -> str:
        """读取文件内容"""
        ext = os.path.splitext(file_path)[1].lower()
        
        # 文本文件
        if ext in ['.txt', '.md', '.py', '.js', '.json', '.csv', '.html', '.xml']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        # PDF文件
        if ext == '.pdf':
            try:
                import pdfplumber
                text_parts = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
                return '\n\n'.join(text_parts)
            except ImportError:
                logger.warning("pdfplumber未安装，尝试使用PyPDF2")
                try:
                    import PyPDF2
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text_parts = [page.extract_text() for page in reader.pages]
                        return '\n\n'.join(filter(None, text_parts))
                except ImportError:
                    raise ImportError("需要安装pdfplumber或PyPDF2来处理PDF文件")
        
        # Word文件
        if ext in ['.docx', '.doc']:
            try:
                from docx import Document
                doc = Document(file_path)
                return '\n\n'.join(para.text for para in doc.paragraphs if para.text)
            except ImportError:
                raise ImportError("需要安装python-docx来处理Word文件")
        
        # 默认按文本读取
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    
    def search(
        self,
        query: str,
        top_k: Optional[int] = None,
        collection_name: Optional[str] = None,
        filter_metadata: Optional[Dict[str, Any]] = None,
        use_rerank: Optional[bool] = None
    ) -> List[RetrievalResult]:
        """
        检索相关内容
        
        Args:
            query: 查询文本
            top_k: 返回数量
            collection_name: 集合名称
            filter_metadata: 元数据过滤
            use_rerank: 是否使用重排序
        
        Returns:
            检索结果列表
        """
        top_k = top_k or self.config.top_k
        use_rerank = use_rerank if use_rerank is not None else self.config.rerank_enabled
        
        # 混合检索
        results = self._retriever.search(
            query=query,
            top_k=top_k * 2 if use_rerank else top_k,  # 重排序需要更多候选
            collection_name=collection_name,
            filter_metadata=filter_metadata,
            use_hybrid=self.config.use_hybrid
        )
        
        # 重排序
        if use_rerank and results:
            results = self._reranker.rerank(
                query=query,
                results=results,
                top_k=top_k
            )
        
        return results
    
    def build_context(
        self,
        query: str,
        max_tokens: int = 4000,
        top_k: int = 5,
        collection_name: Optional[str] = None,
        filter_metadata: Optional[Dict[str, Any]] = None
    ) -> str:
        """
        构建上下文(用于LLM提示)
        
        Args:
            query: 查询文本
            max_tokens: 最大token数
            top_k: 检索数量
            collection_name: 集合名称
            filter_metadata: 元数据过滤
        
        Returns:
            拼接好的上下文文本
        """
        results = self.search(
            query=query,
            top_k=top_k,
            collection_name=collection_name,
            filter_metadata=filter_metadata
        )
        
        if not results:
            return ""
        
        # 按相关性构建上下文，控制总长度
        context_parts = []
        current_tokens = 0
        
        for i, result in enumerate(results):
            # 估算token数
            chunk_tokens = self._estimate_tokens(result.content)
            
            if current_tokens + chunk_tokens > max_tokens:
                break
            
            # 添加来源标注
            source = result.metadata.get('filename', f'文档{i+1}')
            context_parts.append(f"【来源: {source}】\n{result.content}")
            current_tokens += chunk_tokens
        
        return "\n\n---\n\n".join(context_parts)
    
    def _estimate_tokens(self, text: str) -> int:
        """估算token数"""
        chinese_chars = sum(1 for c in text if '\u4e00' <= c <= '\u9fff')
        other_chars = len(text) - chinese_chars
        return int(chinese_chars / 1.5 + other_chars / 4)
    
    def remove_document(
        self,
        doc_id: str,
        collection_name: Optional[str] = None
    ):
        """移除已索引的文档"""
        if doc_id in self._indexed_docs:
            doc_info = self._indexed_docs[doc_id]
            
            # 生成所有块ID
            chunk_ids = [f"{doc_id}_chunk_{i}" for i in range(doc_info.chunk_count)]
            
            # 从检索器移除
            self._retriever.remove_documents(chunk_ids, collection_name)
            
            # 移除记录
            del self._indexed_docs[doc_id]
            
            logger.info(f"文档 {doc_id} 已移除")
    
    def list_indexed_documents(self) -> List[IndexedDocument]:
        """列出所有已索引文档"""
        return list(self._indexed_docs.values())
    
    def get_stats(self) -> Dict[str, Any]:
        """获取RAG引擎统计"""
        return {
            'config': {
                'chunk_size': self.config.chunk_size,
                'chunk_overlap': self.config.chunk_overlap,
                'chunking_strategy': self.config.chunking_strategy.value,
                'embedding_provider': self.config.embedding_provider,
                'use_hybrid': self.config.use_hybrid,
                'rerank_enabled': self.config.rerank_enabled
            },
            'indexed_documents': len(self._indexed_docs),
            'total_chunks': sum(doc.chunk_count for doc in self._indexed_docs.values()),
            'total_tokens': sum(doc.total_tokens for doc in self._indexed_docs.values()),
            'embedding_info': self._embedding_service.get_model_info(),
            'retriever_stats': self._retriever.get_stats()
        }


# 全局RAG引擎实例
_global_rag_engine: Optional[RAGEngine] = None


def get_rag_engine() -> RAGEngine:
    """获取全局RAG引擎"""
    global _global_rag_engine
    if _global_rag_engine is None:
        _global_rag_engine = RAGEngine()
    return _global_rag_engine


def init_rag_engine(
    config: Optional[RAGConfig] = None,
    embedding_api_key: Optional[str] = None,
    llm_client: Optional[Any] = None
) -> RAGEngine:
    """初始化全局RAG引擎"""
    global _global_rag_engine
    _global_rag_engine = RAGEngine(
        config=config,
        embedding_api_key=embedding_api_key,
        llm_client=llm_client
    )
    return _global_rag_engine
