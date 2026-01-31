from flask import Blueprint, request, send_file, jsonify
from docx import Document
import io

export_bp = Blueprint('export', __name__, url_prefix='/api/export')


# ==============================
# docx 导出相关函数
# ==============================

def process_node(doc, node):
    """递归处理 TipTap JSON 节点并写入 docx"""
    node_type = node.get('type')
    content = node.get('content', [])
    
    # 提取纯文本内容
    text = ""
    
    # 根据节点类型处理
    if node_type == 'doc':
        for child in content:
            process_node(doc, child)
            
    elif node_type == 'paragraph':
        # 如果是空的段落，也需要添加
        p = doc.add_paragraph()
        for child in content:
            if child.get('type') == 'text':
                run = p.add_run(child.get('text', ''))
                # 处理加粗、斜体等 marks
                if 'marks' in child:
                    for mark in child['marks']:
                        if mark['type'] == 'bold': run.bold = True
                        if mark['type'] == 'italic': run.italic = True
                        # 可以添加更多样式支持，如下划线等
                        if mark['type'] == 'underline': run.underline = True
                        
    elif node_type == 'heading':
        level = node.get('attrs', {}).get('level', 1)
        # 提取标题文本
        heading_text = ""
        for child in content:
            if child.get('type') == 'text':
                heading_text += child.get('text', '')
        
        doc.add_heading(heading_text, level=level)
        
    elif node_type == 'bulletList':
        for child in content:
            if child.get('type') == 'listItem':
                # docx 库对列表支持比较简单，通常只能通过 paragraph style 控制
                # 这里简化处理，直接用 List Bullet 样式
                # 注意：嵌套列表在这个简化版可能无法完美展现
                p = doc.add_paragraph(style='List Bullet')
                # 列表项里面通常是 paragraph
                for item_content in child.get('content', []):
                    if item_content.get('type') == 'paragraph':
                        for text_node in item_content.get('content', []):
                            if text_node.get('type') == 'text':
                                run = p.add_run(text_node.get('text', ''))
                                if 'marks' in text_node:
                                    for mark in text_node['marks']:
                                        if mark['type'] == 'bold': run.bold = True
                                        if mark['type'] == 'italic': run.italic = True
                                        
    elif node_type == 'orderedList':
        for child in content:
            if child.get('type') == 'listItem':
                p = doc.add_paragraph(style='List Number')
                for item_content in child.get('content', []):
                    if item_content.get('type') == 'paragraph':
                        for text_node in item_content.get('content', []):
                            if text_node.get('type') == 'text':
                                run = p.add_run(text_node.get('text', ''))
                                if 'marks' in text_node:
                                    for mark in text_node['marks']:
                                        if mark['type'] == 'bold': run.bold = True
                                        if mark['type'] == 'italic': run.italic = True

    elif node_type == 'codeBlock':
       # 代码块用特定样式或单倍行距+固定字体
       # 暂时简化为普通文本
        code_text = ""
        for child in content:
            if child.get('type') == 'text':
               code_text += child.get('text', '')
        
        if code_text:
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing' # 紧凑样式

    elif node_type == 'blockquote':
        # 引用块
         for child in content:
            if child.get('type') == 'paragraph':
                text = ""
                for text_node in child.get('content', []):
                    text += text_node.get('text', '')
                if text:
                    p = doc.add_paragraph(text, style='Quote')


@export_bp.route('/docx', methods=['POST'])
def export_docx():
    try:
        data = request.json
        if not data:
            return jsonify({'error': 'No data provided'}), 400
            
        title = data.get('title', 'Document')
        editor_json = data.get('content', {})

        # 创建 Word 文档
        doc = Document()
        doc.add_heading(title, 0)

        # 递归处理内容
        if not editor_json or not isinstance(editor_json, dict):
             process_node(doc, {'type': 'doc', 'content': []}) # TODO_202 应该返回报错
        else:
             process_node(doc, editor_json)

        # 保存到内存流
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        return send_file(
            file_stream,
            as_attachment=True,
            download_name=f"{title}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        print(f"Export error: {e}")
        return jsonify({'error': str(e)}), 500
